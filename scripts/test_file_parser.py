"""
Test script for src/file_parser.py

Tests text/md parsing, empty file handling, and optional docx parsing.
Does NOT use real user resumes.
"""

import sys
import os
import unittest
from io import BytesIO

# Add project root to sys.path so "from src.xxx" works when run directly
_PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

# ---------------------------------------------------------------------------
# Fake UploadedFile helper
# ---------------------------------------------------------------------------

class FakeUpload(BytesIO):
    """Minimal streamlit UploadedFile stand-in for testing."""
    def __init__(self, name: str, data: bytes):
        super().__init__(data)
        self.name = name


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestFileParser(unittest.TestCase):

    def test_txt_utf8(self):
        from src.file_parser import extract_text_from_upload
        content = "张三\n计算机科学\n技能：Python、PyTorch\n项目：LLM 求职助手\n"
        f = FakeUpload("resume.txt", content.encode("utf-8"))
        text, msg = extract_text_from_upload(f)
        self.assertIn("张三", text)
        self.assertIn("Python", text)
        self.assertTrue(msg.startswith("✅"))

    def test_txt_gbk(self):
        """Test GBK/GB18030 encoded text file."""
        from src.file_parser import extract_text_from_upload
        # Use a longer string to pass the 30-char threshold
        content = (
            "李四\n"
            "算法工程师\n"
            "技能:C++,CUDA,Python,PyTorch\n"
            "项目:推荐系统开发、召回排序优化\n"
            "教育:计算机科学硕士\n"
        )
        try:
            data = content.encode("gb18030")
        except UnicodeEncodeError:
            self.skipTest("无法将测试内容编码为 GB18030")
        f = FakeUpload("resume_gbk.txt", data)
        text, msg = extract_text_from_upload(f)
        self.assertIn("李四", text, f"GB18030 解码失败, msg={msg}")
        self.assertTrue(msg.startswith("✅"), f"期望成功消息，实际: {msg}")

    def test_md(self):
        from src.file_parser import extract_text_from_upload
        content = "# 张同学\n## 技能\n- Python\n- PyTorch\n## 项目\nLLM 求职助手\n"
        f = FakeUpload("resume.md", content.encode("utf-8"))
        text, msg = extract_text_from_upload(f)
        self.assertIn("张同学", text)
        self.assertTrue(msg.startswith("✅"))

    def test_empty_file(self):
        from src.file_parser import extract_text_from_upload
        f = FakeUpload("empty.txt", b"")
        text, msg = extract_text_from_upload(f)
        self.assertEqual(text, "")
        self.assertIn("文件为空", msg)  # empty file

    def test_too_short(self):
        from src.file_parser import extract_text_from_upload
        f = FakeUpload("short.txt", "Hi".encode("utf-8"))
        text, msg = extract_text_from_upload(f)
        self.assertEqual(text, "")
        self.assertIn("过少", msg)

    def test_doc_not_supported(self):
        from src.file_parser import extract_text_from_upload
        f = FakeUpload("resume.doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
        text, msg = extract_text_from_upload(f)
        self.assertEqual(text, "")
        self.assertIn(".doc", msg)
        self.assertIn("另存为", msg)

    def test_docx_import_failure_message(self):
        """If python-docx is not installed, the error message should guide the user."""
        from src.file_parser import _read_docx
        # We call _read_docx directly with a fake file;
        # if docx is not installed we should get a helpful message.
        f = FakeUpload("resume.docx", b"PK\x03\x04fake zip content")
        # This will fail at import time inside _read_docx
        text, msg = _read_docx(f, "resume.docx")
        # Either parsed (if docx installed) or got a clear error message
        if not text:
            self.assertIn("未安装", msg)  # should mention missing dependency

    def test_pdf_import_failure_message(self):
        """If pypdf is not installed, the error message should guide the user."""
        from src.file_parser import _read_pdf
        f = FakeUpload("resume.pdf", b"%PDF-1.4 fake pdf content")
        text, msg = _read_pdf(f, "resume.pdf")
        if not text:
            self.assertIn("未安装", msg)


# ---------------------------------------------------------------------------
# Optional: generate a real .docx and test parsing
# ---------------------------------------------------------------------------

def _test_docx_real():
    """Generate a temporary .docx and verify parsing."""
    try:
        import docx
    except ImportError:
        print("⚠️  python-docx not installed, skipping real .docx test")
        return

    from src.file_parser import extract_text_from_upload
    from io import BytesIO as _BytesIO
    import tempfile, os

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc = docx.Document()
        doc.add_paragraph("王同学")
        doc.add_paragraph("岗位：推荐算法工程师")
        doc.add_paragraph("技能：Python, PyTorch, Transformer")
        doc.save(tmp_path)

        try:
            data = open(tmp_path, "rb").read()
            f = FakeUpload("test_real.docx", data)
            text, msg = extract_text_from_upload(f)
            print(f"[docx real] msg={msg}")
            assert "王同学" in text, "docx parse should contain name"
            print("✅  real .docx parsing OK")
        finally:
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("=" * 60)
    print("  File Parser Tests")
    print("=" * 60)

    # Run unittest
    suite = unittest.TestLoader().loadTestsFromTestCase(TestFileParser)
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    # Optional real .docx test
    print()
    _test_docx_real()

    sys.exit(0 if result.wasSuccessful() else 1)
