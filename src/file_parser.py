"""
File parser for Offer Catcher.

Supports: .txt, .md, .pdf, .docx
Pure-local parsing — no file saved to disk, no API calls.
"""

from io import BytesIO
from typing import Tuple


def extract_text_from_upload(uploaded_file) -> Tuple[str, str]:
    """
    Read an uploaded file (Streamlit UploadedFile) and extract text.

    Returns:
        (text, message)
        - text: extracted resume text, or "" on failure
        - message: UI-friendly status message (success / error / warning)
    """
    if uploaded_file is None:
        return "", "未上传文件。"

    filename = getattr(uploaded_file, "name", "unknown")
    suffix = ""
    if "." in filename:
        suffix = "." + filename.rsplit(".", 1)[-1].lower()

    # --- .txt / .md ---
    if suffix in (".txt", ".md"):
        return _read_text_file(uploaded_file, filename)

    # --- .pdf ---
    if suffix == ".pdf":
        return _read_pdf(uploaded_file, filename)

    # --- .docx ---
    if suffix == ".docx":
        return _read_docx(uploaded_file, filename)

    # --- .doc (not supported) ---
    if suffix == ".doc":
        return (
            "",
            f"❌ 不支持 .doc 格式（{filename}）。\n"
            f"请先将文件另存为 .docx 或 PDF，再重新上传。",
        )

    return "", f"❌ 不支持的文件格式：{suffix}（{filename}）"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _read_text_file(uploaded_file, filename: str) -> Tuple[str, str]:
    """Read .txt / .md as text, trying multiple encodings."""
    raw = uploaded_file.read()
    # Reset pointer so downstream calls can re-read if needed
    if hasattr(uploaded_file, "seek"):
        try:
            uploaded_file.seek(0)
        except Exception:
            pass

    if not raw:
        return "", f"❌ 文件为空（{filename}）。"

    encodings = ["utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"]
    text = None
    for enc in encodings:
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        return "", f"❌ 无法识别文件编码（{filename}）。请另存为 UTF-8 文本后重试。"

    text = text.strip()
    if len(text) < 30:
        return "", f"⚠️ 文件内容过少（仅 {len(text)} 字符，{filename}），请检查文件是否正确。"

    return text, f"✅ 已解析 {filename}（{len(text)} 字符，文本格式）"


def _read_pdf(uploaded_file, filename: str) -> Tuple[str, str]:
    """Extract text from PDF using pypdf."""
    try:
        import pypdf
    except ImportError:
        return (
            "",
            f"❌ 服务器未安装 pypdf，无法解析 PDF（{filename}）。\n"
            f"请在 requirements.txt 中添加 pypdf>=4.0.0 并重新部署。",
        )

    try:
        uploaded_file.seek(0)
        reader = pypdf.PdfReader(uploaded_file)
        parts = []
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            except Exception:
                continue
        text = "\n".join(parts).strip()

        if len(text) < 30:
            return (
                "",
                f"⚠️ PDF 解析结果过少（仅 {len(text)} 字符，{filename}）。\n"
                f"可能是扫描版 PDF，请尝试复制文本后粘贴，或上传 DOCX/PDF 文字版。",
            )

        return text, f"✅ 已解析 {filename}（{len(text)} 字符，PDF {len(reader.pages)} 页）"
    except Exception as e:
        return "", f"❌ PDF 解析失败（{filename}）：{e}"


def _read_docx(uploaded_file, filename: str) -> Tuple[str, str]:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
    except ImportError:
        return (
            "",
            f"❌ 服务器未安装 python-docx，无法解析 DOCX（{filename}）。\n"
            f"请在 requirements.txt 中添加 python-docx>=1.1.0 并重新部署。",
        )

    try:
        uploaded_file.seek(0)
        doc = docx.Document(uploaded_file)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        row_texts.append(t)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        text = "\n".join(parts).strip()

        if len(text) < 30:
            return (
                "",
                f"⚠️ DOCX 解析结果过少（仅 {len(text)} 字符，{filename}），请检查文件是否正确。",
            )

        return text, f"✅ 已解析 {filename}（{len(text)} 字符，DOCX）"
    except Exception as e:
        return "", f"❌ DOCX 解析失败（{filename}）：{e}"
